import json
import os
import re
import uuid
from typing import List, Optional

import chromadb
import fitz  # PyMuPDF
import numpy as np
import requests
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from PIL import Image
from pydantic import BaseModel
from pypdf import PdfReader
from rapidocr_onnxruntime import RapidOCR
from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer

app = FastAPI(title="AI Ingestion Backend")

# =========================
# CONFIG
# =========================
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/var/data/uploads")
CHROMA_DIR = os.getenv("CHROMA_DIR", "/var/data/chroma")
OLLAMA_URL = os.getenv("OLLAMA_URL", "").strip()
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "gemma3:4b")
ROUTER_MODEL = os.getenv("ROUTER_MODEL", "modelfilerouter")
ENABLE_OCR = os.getenv("ENABLE_OCR", "true").lower() == "true"

EMBEDDING_MODEL_NAME = os.getenv(
    "EMBEDDING_MODEL_NAME",
    "sentence-transformers/all-MiniLM-L6-v2",
)
TOKEN_CHUNK_SIZE = int(os.getenv("TOKEN_CHUNK_SIZE", "220"))
TOKEN_CHUNK_OVERLAP = int(os.getenv("TOKEN_CHUNK_OVERLAP", "40"))
TOP_K = int(os.getenv("TOP_K", "3"))
OCR_DPI = int(os.getenv("OCR_DPI", "180"))

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(CHROMA_DIR, exist_ok=True)

# =========================
# MODELS / CLIENTS
# =========================
embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
tokenizer = AutoTokenizer.from_pretrained(EMBEDDING_MODEL_NAME)
chroma_client = chromadb.PersistentClient(path=CHROMA_DIR)
ocr_engine = RapidOCR()

ALLOWED_TOOLS = {
    "employees": ["retrieval"],
    "create_service": ["mutation"],
}

# =========================
# REQUEST MODELS
# =========================
class SearchRequest(BaseModel):
    user_id: str
    query: str
    top_k: int = TOP_K


class AskAIRequest(BaseModel):
    user_id: str
    question: str
    top_k: int = TOP_K


class RouteRequest(BaseModel):
    message: str


# =========================
# HELPERS
# =========================
def sanitize_user_id(user_id: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_-]", "_", user_id.strip())
    if not cleaned:
        raise HTTPException(status_code=400, detail="Invalid user_id.")
    return cleaned


def safe_collection_name(user_id: str) -> str:
    return f"documents_{sanitize_user_id(user_id)}"


def get_user_collection(user_id: str):
    return chroma_client.get_or_create_collection(name=safe_collection_name(user_id))


def get_user_folder(user_id: str) -> str:
    folder = os.path.join(UPLOAD_DIR, sanitize_user_id(user_id))
    os.makedirs(folder, exist_ok=True)
    return folder


def normalize_whitespace(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    full_text: List[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            full_text.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                full_text.append(" | ".join(row_text))

    return normalize_whitespace("\n\n".join(full_text))


def extract_text_from_pdf_native(file_path: str) -> str:
    reader = PdfReader(file_path)
    parts: List[str] = []

    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(text)

    return normalize_whitespace("\n\n".join(parts))


def ocr_pdf_with_pymupdf(file_path: str, dpi: int = OCR_DPI) -> str:
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    doc = fitz.open(file_path)
    page_texts: List[str] = []

    try:
        for page in doc:
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            mode = "RGB" if pix.n < 4 else "RGBA"
            image = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
            image_np = np.array(image)

            ocr_result, _ = ocr_engine(image_np)
            if not ocr_result:
                continue

            lines: List[str] = []
            for item in ocr_result:
                if isinstance(item, (list, tuple)) and len(item) >= 2:
                    line_text = str(item[1]).strip()
                    if line_text:
                        lines.append(line_text)

            if lines:
                page_texts.append("\n".join(lines))
    finally:
        doc.close()

    return normalize_whitespace("\n\n".join(page_texts))


def extract_text_from_pdf(file_path: str) -> str:
    native_text = extract_text_from_pdf_native(file_path)

    if len(native_text) >= 100:
        return native_text

    if not ENABLE_OCR:
        if native_text:
            return native_text
        raise HTTPException(
            status_code=400,
            detail="No usable text could be extracted from the PDF. OCR is disabled.",
        )

    ocr_text = ocr_pdf_with_pymupdf(file_path)

    if ocr_text:
        combined = f"{native_text}\n\n{ocr_text}".strip() if native_text else ocr_text
        return normalize_whitespace(combined)

    if native_text:
        return native_text

    raise HTTPException(
        status_code=400,
        detail="No text could be extracted from the PDF, even after OCR fallback.",
    )


def chunk_text_by_tokens(
    text: str,
    chunk_size: int = TOKEN_CHUNK_SIZE,
    overlap: int = TOKEN_CHUNK_OVERLAP,
) -> List[str]:
    text = normalize_whitespace(text)
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: List[str] = []

    current_token_ids: List[int] = []
    current_text_parts: List[str] = []

    def flush_current():
        nonlocal current_token_ids, current_text_parts
        if current_text_parts:
            chunk = normalize_whitespace("\n\n".join(current_text_parts))
            if chunk:
                chunks.append(chunk)

    for paragraph in paragraphs:
        para_ids = tokenizer.encode(paragraph, add_special_tokens=False)

        if len(para_ids) > chunk_size:
            flush_current()
            current_token_ids = []
            current_text_parts = []

            start = 0
            while start < len(para_ids):
                end = min(start + chunk_size, len(para_ids))
                piece_ids = para_ids[start:end]
                piece_text = tokenizer.decode(piece_ids, skip_special_tokens=True).strip()
                if piece_text:
                    chunks.append(normalize_whitespace(piece_text))
                if end == len(para_ids):
                    break
                start = max(end - overlap, start + 1)
            continue

        projected_len = len(current_token_ids) + len(para_ids)

        if projected_len <= chunk_size:
            current_text_parts.append(paragraph)
            current_token_ids.extend(para_ids)
        else:
            flush_current()

            if overlap > 0 and current_token_ids:
                overlap_ids = current_token_ids[-overlap:]
                overlap_text = tokenizer.decode(overlap_ids, skip_special_tokens=True).strip()
                current_text_parts = [overlap_text] if overlap_text else []
                current_token_ids = overlap_ids[:]
            else:
                current_text_parts = []
                current_token_ids = []

            current_text_parts.append(paragraph)
            current_token_ids.extend(para_ids)

    flush_current()

    cleaned_chunks: List[str] = []
    seen = set()
    for chunk in chunks:
        normalized = chunk.strip()
        if len(normalized) < 20:
            continue
        if normalized not in seen:
            seen.add(normalized)
            cleaned_chunks.append(normalized)

    return cleaned_chunks


def generate_embeddings(chunks: List[str]) -> List[List[float]]:
    return embedding_model.encode(chunks).tolist()


def delete_existing_file_chunks(collection, filename: str):
    try:
        existing = collection.get(where={"filename": filename})
        existing_ids = existing.get("ids", [])
        if existing_ids:
            collection.delete(ids=existing_ids)
    except Exception:
        pass


def store_in_chroma(collection, filename: str, user_id: str, chunks: List[str], embeddings: List[List[float]]):
    ids = []
    metadatas = []

    for i, _chunk in enumerate(chunks):
        ids.append(str(uuid.uuid4()))
        metadatas.append(
            {
                "user_id": sanitize_user_id(user_id),
                "filename": filename,
                "chunk_index": i,
            }
        )

    collection.add(
        documents=chunks,
        embeddings=embeddings,
        metadatas=metadatas,
        ids=ids,
    )


def extract_first_json(text: str) -> Optional[str]:
    match = re.search(r"\{.*?\}", text, re.DOTALL)
    return match.group(0) if match else None


def is_valid_router_output(parsed: dict) -> bool:
    if not isinstance(parsed, dict):
        return False

    route = parsed.get("route")

    if route == "text":
        response = parsed.get("response")
        return isinstance(response, str) and response.strip() != ""

    if route == "tools":
        toolname = parsed.get("toolname")
        mode = parsed.get("mode")

        if toolname not in ALLOWED_TOOLS:
            return False

        return mode in ALLOWED_TOOLS[toolname]

    return False


def ask_ollama(prompt: str, model_name: Optional[str] = None) -> str:
    if not OLLAMA_URL:
        raise HTTPException(
            status_code=503,
            detail="OLLAMA_URL is not configured. Set it in environment variables.",
        )

    selected_model = model_name or OLLAMA_MODEL

    try:
        response = requests.post(
            f"{OLLAMA_URL.rstrip('/')}/api/generate",
            json={
                "model": selected_model,
                "prompt": prompt,
                "stream": False,
            },
            timeout=180,
        )
        response.raise_for_status()
        return response.json().get("response", "No response returned from Ollama.")
    except requests.RequestException as e:
        raise HTTPException(
            status_code=500,
            detail=f"Ollama request failed: {str(e)}",
        )


# =========================
# ROUTES
# =========================
@app.get("/")
def home():
    return {"message": "AI Ingestion Backend Running"}


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/upload-document")
async def upload_document(
    user_id: str = Form(...),
    replace_existing: bool = Form(True),
    file: UploadFile = File(...),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is missing.")

    lower_name = file.filename.lower()
    if not (lower_name.endswith(".docx") or lower_name.endswith(".pdf")):
        raise HTTPException(
            status_code=400,
            detail="Only .docx and .pdf files are supported in this version.",
        )

    user_folder = get_user_folder(user_id)
    file_path = os.path.join(user_folder, file.filename)

    with open(file_path, "wb") as f:
        f.write(await file.read())

    if lower_name.endswith(".docx"):
        extracted_text = extract_text_from_docx(file_path)
    else:
        extracted_text = extract_text_from_pdf(file_path)

    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="No text could be extracted from the uploaded file.")

    chunks = chunk_text_by_tokens(extracted_text)

    if not chunks:
        raise HTTPException(status_code=400, detail="No usable chunks were created from the uploaded file.")

    embeddings = generate_embeddings(chunks)
    collection = get_user_collection(user_id)

    if replace_existing:
        delete_existing_file_chunks(collection, file.filename)

    store_in_chroma(collection, file.filename, user_id, chunks, embeddings)

    return {
        "user_id": sanitize_user_id(user_id),
        "filename": file.filename,
        "message": "File uploaded, text extracted, chunked, embedded, and stored successfully.",
        "path": file_path,
        "chunk_count": len(chunks),
        "embedding_count": len(embeddings),
        "stored_in_vector_db": True,
        "first_chunk_preview": chunks[0][:500],
    }


@app.post("/search")
async def search_documents(request: SearchRequest):
    collection = get_user_collection(request.user_id)
    query_embedding = embedding_model.encode([request.query]).tolist()[0]

    results = collection.query(
        query_embeddings=[query_embedding],
        n_results=max(1, request.top_k),
    )

    return {
        "user_id": sanitize_user_id(request.user_id),
        "query": request.query,
        "results": results,
    }


@app.post("/route-message")
async def route_message(request: RouteRequest):
    raw_response = ask_ollama(request.message, model_name=ROUTER_MODEL)

    cleaned = extract_first_json(raw_response)
    fallback = {
        "route": "text",
        "response": "Hello! How can I help you?"
    }

    if not cleaned:
        return {
            "input": request.message,
            "raw_response": raw_response,
            "route_result": fallback,
            "fallback_used": True,
        }

    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError:
        return {
            "input": request.message,
            "raw_response": raw_response,
            "route_result": fallback,
            "fallback_used": True,
        }

    if not is_valid_router_output(parsed):
        return {
            "input": request.message,
            "raw_response": raw_response,
            "route_result": fallback,
            "fallback_used": True,
        }

    return {
        "input": request.message,
        "raw_response": raw_response,
        "route_result": parsed,
        "fallback_used": False,
    }


@app.post("/ask-ai")
async def ask_ai(request: AskAIRequest):
    collection = get_user_collection(request.user_id)
    query_embedding = embedding_model.encode([request.question]).tolist()[0]

    results = collection.query(
        query_embeddings=[query_embedding],
        n_results=max(1, request.top_k),
    )

    documents = results.get("documents", [])
    metadatas = results.get("metadatas", [])

    matched_chunks = documents[0] if documents and len(documents) > 0 else []
    matched_metadata = metadatas[0] if metadatas and len(metadatas) > 0 else []

    if not matched_chunks:
        return {
            "user_id": sanitize_user_id(request.user_id),
            "question": request.question,
            "answer": "The answer was not found in the uploaded document.",
            "source_chunks": [],
            "source_metadata": [],
        }

    context = "\n\n".join(matched_chunks)

    prompt = f"""
You are answering based only on the document context below.

Document Context:
{context}

Question:
{request.question}

Instructions:
- Answer using only the document context.
- If the answer is not in the context, say exactly: "The answer was not found in the uploaded document."
"""

    answer = ask_ollama(prompt)

    return {
        "user_id": sanitize_user_id(request.user_id),
        "question": request.question,
        "answer": answer,
        "source_chunks": matched_chunks,
        "source_metadata": matched_metadata,
    }
